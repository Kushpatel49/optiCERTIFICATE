import streamlit as st
import datetime
from dataclasses import dataclass, field
from typing import List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import io

# CA Partner Details
CA_PARTNERS = {
    "CA HARSH B PATEL": {"membership_no": "600794"},
    "CA PRERIT PAREKH": {"membership_no": "194438"}
}

# ==================== DATA MODELS ====================

@dataclass
class BankAccount:
    holder_name: str
    account_number: str
    bank_name: str
    balance_inr: float
    statement_date: str = ""
    
    @property
    def balance_foreign(self) -> float:
        return self.balance_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class InsurancePolicy:
    holder_name: str
    policy_number: str
    amount_inr: float
    
    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class PFAccount:
    holder_name: str
    pf_account_number: str
    amount_inr: float
    
    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class Deposit:
    holder_name: str
    account_number: str
    amount_inr: float
    
    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class NPSAccount:
    owner_name: str
    pran_number: str
    amount_inr: float

    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class MutualFund:
    holder_name: str
    folio_number: str
    policy_name: str
    amount_inr: float

    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class Share:
    company_name: str
    num_shares: int
    market_price_inr: float

    @property
    def amount_inr(self) -> float:
        return self.num_shares * self.market_price_inr

    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class Vehicle:
    vehicle_type: str
    make_model_year: str
    registration_number: str
    market_value_inr: float

    @property
    def amount_foreign(self) -> float:
        return self.market_value_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class PostOfficeScheme:
    scheme_type: str
    account_number: str
    amount_inr: float
    
    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class GoldHolding:
    owner_name: str
    weight_grams: float
    rate_per_10g: float
    valuation_date: str = ""
    valuer_name: str = ""
    
    @property
    def amount_inr(self) -> float:
        return (self.weight_grams / 10) * self.rate_per_10g
    
    @property
    def amount_foreign(self) -> float:
        return self.amount_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class Property:
    owner_name: str
    property_type: str
    address: str
    valuation_inr: float
    valuation_date: str = ""
    valuer_name: str = ""
    
    @property
    def valuation_foreign(self) -> float:
        return self.valuation_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class Liability:
    description: str
    amount_inr: float
    details: str = ""

@dataclass
class PartnershipFirm:
    firm_name: str
    partner_name: str
    holding_percentage: float
    capital_balance_inr: float
    valuation_date: str

    @property
    def amount_foreign(self) -> float:
        return self.capital_balance_inr / st.session_state.get('exchange_rate', 63.34)

@dataclass
class NetWorthData:
    # Personal Details
    individual_name: str
    individual_address: str
    certificate_date: str
    engagement_date: str
    embassy_name: str
    embassy_address: str
    passport_number: str = ""
    foreign_currency: str = "CAD"
    exchange_rate: float = 63.34
    
    # Assets
    bank_accounts: List[BankAccount] = field(default_factory=list)
    insurance_policies: List[InsurancePolicy] = field(default_factory=list)
    pf_accounts: List[PFAccount] = field(default_factory=list)
    deposits: List[Deposit] = field(default_factory=list)
    nps_accounts: List[NPSAccount] = field(default_factory=list)
    mutual_funds: List[MutualFund] = field(default_factory=list)
    shares: List[Share] = field(default_factory=list)
    vehicles: List[Vehicle] = field(default_factory=list)
    post_office_schemes: List[PostOfficeScheme] = field(default_factory=list)
    partnership_firms: List[PartnershipFirm] = field(default_factory=list)
    gold_holdings: List[GoldHolding] = field(default_factory=list)
    properties: List[Property] = field(default_factory=list)
    
    # Liabilities
    liabilities: List[Liability] = field(default_factory=list)
    
    # CA Details (Pre-filled)
    ca_firm_name: str = "Patel Parekh & Associates"
    ca_frn: str = "154335W"
    ca_partner_name: str = "CA HARSH B PATEL"
    ca_membership_no: str = "600794"
    ca_designation: str = "Partner"
    ca_place: str = "Vijapur"
    
    @property
    def total_bank_balance_inr(self) -> float:
        return sum(acc.balance_inr for acc in self.bank_accounts)
    
    @property
    def total_bank_balance_foreign(self) -> float:
        return self.total_bank_balance_inr / self.exchange_rate
    
    @property
    def total_insurance_inr(self) -> float:
        return sum(pol.amount_inr for pol in self.insurance_policies)
    
    @property
    def total_insurance_foreign(self) -> float:
        return self.total_insurance_inr / self.exchange_rate
    
    @property
    def total_pf_accounts_inr(self) -> float:
        return sum(acc.amount_inr for acc in self.pf_accounts)
    
    @property
    def total_pf_accounts_foreign(self) -> float:
        return self.total_pf_accounts_inr / self.exchange_rate

    @property
    def total_deposits_inr(self) -> float:
        return sum(dep.amount_inr for dep in self.deposits)
    
    @property
    def total_deposits_foreign(self) -> float:
        return self.total_deposits_inr / self.exchange_rate
    
    @property
    def total_nps_inr(self) -> float:
        return sum(nps.amount_inr for nps in self.nps_accounts)

    @property
    def total_nps_foreign(self) -> float:
        return self.total_nps_inr / self.exchange_rate

    @property
    def total_mutual_funds_inr(self) -> float:
        return sum(mf.amount_inr for mf in self.mutual_funds)

    @property
    def total_mutual_funds_foreign(self) -> float:
        return self.total_mutual_funds_inr / self.exchange_rate

    @property
    def total_shares_inr(self) -> float:
        return sum(s.amount_inr for s in self.shares)

    @property
    def total_shares_foreign(self) -> float:
        return self.total_shares_inr / self.exchange_rate

    @property
    def total_vehicles_inr(self) -> float:
        return sum(v.market_value_inr for v in self.vehicles)

    @property
    def total_vehicles_foreign(self) -> float:
        return self.total_vehicles_inr / self.exchange_rate

    @property
    def total_post_office_inr(self) -> float:
        return sum(p.amount_inr for p in self.post_office_schemes)

    @property
    def total_post_office_foreign(self) -> float:
        return self.total_post_office_inr / self.exchange_rate
    
    @property
    def total_partnership_firms_inr(self) -> float:
        return sum(f.capital_balance_inr for f in self.partnership_firms)

    @property
    def total_partnership_firms_foreign(self) -> float:
        return self.total_partnership_firms_inr / self.exchange_rate
    
    @property
    def total_gold_inr(self) -> float:
        return sum(gold.amount_inr for gold in self.gold_holdings)
    
    @property
    def total_gold_foreign(self) -> float:
        return self.total_gold_inr / self.exchange_rate
    
    @property
    def total_movable_assets_inr(self) -> float:
        return (self.total_bank_balance_inr + 
                self.total_insurance_inr + 
                self.total_pf_accounts_inr + 
                self.total_deposits_inr +
                self.total_nps_inr +
                self.total_mutual_funds_inr +
                self.total_shares_inr +
                self.total_vehicles_inr +
                self.total_post_office_inr +
                self.total_partnership_firms_inr +
                self.total_gold_inr)
    
    @property
    def total_movable_assets_foreign(self) -> float:
        return self.total_movable_assets_inr / self.exchange_rate
    
    @property
    def total_immovable_assets_inr(self) -> float:
        return sum(prop.valuation_inr for prop in self.properties)
    
    @property
    def total_immovable_assets_foreign(self) -> float:
        return self.total_immovable_assets_inr / self.exchange_rate
    
    @property
    def total_liabilities_inr(self) -> float:
        return sum(liab.amount_inr for liab in self.liabilities)
    
    @property
    def net_worth_inr(self) -> float:
        return self.total_movable_assets_inr + self.total_immovable_assets_inr - self.total_liabilities_inr
    
    @property
    def net_worth_foreign(self) -> float:
        return self.net_worth_inr / self.exchange_rate

@dataclass
class PartnershipFirm:
    firm_name: str
    partner_name: str
    holding_percentage: float
    capital_balance_inr: float
    valuation_date: str

    @property
    def amount_foreign(self) -> float:
        return self.capital_balance_inr / st.session_state.get('exchange_rate', 63.34)

# ==================== DOCUMENT GENERATOR ====================

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'tc{}'.format(edge.capitalize())
            element = OxmlElement('w:{}'.format(tag))
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)

def add_table_with_borders(doc, rows, cols):
    """Add table with borders"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = True  # Set layout to autofit content
    return table

def generate_networth_certificate(data: NetWorthData) -> Document:
    """Generate the complete Net Worth Certificate document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Independent Practitioner's Certificate on Net Worth where no Books of")
    title_run.bold = True
    title_run.underline = True
    title_run.add_break()
    title_run = title.add_run("Account have been maintained (For VISA Application Purpose)")
    title_run.bold = True
    title_run.underline = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # To Address
    to_para = doc.add_paragraph('To')
    to_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    embassy_para = doc.add_paragraph(data.embassy_name)
    embassy_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    embassy_lines = data.embassy_address.split('\n')
    for line in embassy_lines:
        line_para = doc.add_paragraph(line)
        line_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Certificate Body
    para = doc.add_paragraph()
    run = para.add_run("Independent Practitioner's Certificate on Net Worth where no Books of")
    run.underline = True
    run.add_break()
    run = para.add_run("Account have been maintained (For VISA Application Purpose)")
    run.underline = True
    
    cert_para1 = doc.add_paragraph(f'1. This Certificate is issued in accordance with the terms of my/our engagement letter/agreement dated {data.engagement_date}.')
    cert_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    name_with_passport = f"{data.individual_name} (Passport No.: {data.passport_number})" if data.passport_number else data.individual_name
    cert_para2 = doc.add_paragraph(f'2. I/we have been engaged by Mr./Ms. {name_with_passport} (hereinafter referred to as the "individual") having residential address at {data.individual_address} to certify the Net Worth as at {data.certificate_date} for submission to {data.embassy_name} for VISA application purpose.')
    cert_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Individual's Responsibility
    para = doc.add_paragraph()
    run = para.add_run("Individual's Responsibility")
    run.bold = True
    run.underline = True
    
    resp_para = doc.add_paragraph(f'3. The preparation and presentation of the Statement of Net Worth ("the Statement") as at {data.certificate_date} are the responsibility of the individual, including the preparation and maintenance of all relevant supporting records and documents. This responsibility includes the design, implementation and maintenance of adequate internal controls relevant to the accuracy and completeness of the Statement.')
    resp_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Practitioner's Responsibility
    para = doc.add_paragraph()
    run = para.add_run("Practitioner's Responsibility")
    run.bold = True
    run.underline = True
    
    prac_para = doc.add_paragraph(f'4. It is my/our responsibility to examine the Statement of Net Worth prepared by the individual as at {data.certificate_date} and to certify whether the same is based on relevant supporting records and documents as made available to us.')
    prac_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    guide_para = doc.add_paragraph('5. I/we conducted my/our examination of the Statement in accordance with the Guidance Note on Reports or Certificates for Special Purposes (Revised 2016) issued by the Institute of Chartered Accountants of India ("ICAI"). The Guidance Note requires that I/we comply with the ethical requirements of the Code of Ethics issued by ICAI.')
    guide_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    sqc_para = doc.add_paragraph('6. I/we have complied with the relevant applicable requirements of the Standard on Quality Control (SQC) 1, Quality Control for Firms that Perform Audits and Reviews of Historical Financial Information, and Other Assurance and Related Services Engagements.')
    sqc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Opinion
    para = doc.add_paragraph()
    run = para.add_run("Opinion")
    run.bold = True
    run.underline = True
    
    name_with_passport = f"{data.individual_name} (Passport No.: {data.passport_number})" if data.passport_number else data.individual_name
    opinion_para = doc.add_paragraph(f'7. Based on my/our examination as above and the information and explanations given to me/us, I/we certify that the Statement of Net Worth of Mr./Ms. {name_with_passport} as at {data.certificate_date} annexed herewith shows a Net Worth of Rs. {data.net_worth_inr:,.2f} (Rupees {convert_to_words(data.net_worth_inr)} only) which has been computed on the basis of the details, representations, and documents made available to me/us by the individual.')
    opinion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Restriction on Use
    para = doc.add_paragraph()
    run = para.add_run("Restriction on Use")
    run.bold = True
    run.underline = True
    
    restrict_para = doc.add_paragraph(f'8. This Certificate has been issued at the request of the individual for submission to {data.embassy_name} for VISA application purpose. This Certificate should not be used for any other purpose or by any person other than the addressee of this Certificate. Accordingly, I/we do not accept or assume any liability or any duty of care for any other purpose or to any other person to whom this Certificate is shown or into whose hands it may come without my/our prior consent in writing.')
    restrict_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature Block
    firm_para = doc.add_paragraph(f'FOR {data.ca_firm_name.upper()}')
    firm_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ca_para = doc.add_paragraph('CHARTERED ACCOUNTANTS')
    ca_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    frn_para = doc.add_paragraph(f"FRN: {data.ca_frn}")
    frn_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = True

    # Make table borders invisible using XML
    for row in sig_table.rows:
        for cell in row.cells:
            # Set all borders to nil (invisible)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                  r'<w:top w:val="nil"/>'
                                  r'<w:left w:val="nil"/>'
                                  r'<w:bottom w:val="nil"/>'
                                  r'<w:right w:val="nil"/>'
                                  r'</w:tcBorders>')
            tcPr.append(tcBorders)

    # Left column - CA details
    sig_table.rows[0].cells[0].text = f'{data.ca_partner_name.upper()}'
    sig_table.rows[1].cells[0].text = f'{data.ca_designation.upper()}'
    sig_table.rows[2].cells[0].text = f'MEMBERSHIP NO.: {data.ca_membership_no}'
    sig_table.rows[3].cells[0].text = 'UDIN: [TO BE GENERATED]'

    # Right column - Date and Place
    sig_table.rows[0].cells[1].text = f'DATE: {data.certificate_date}'
    sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sig_table.rows[1].cells[1].text = f'PLACE: {data.ca_place.upper()}'
    sig_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Leave bottom two cells in right column empty for spacing
    sig_table.rows[2].cells[1].text = ''
    sig_table.rows[3].cells[1].text = ''
    
    doc.add_paragraph()

    name_with_passport = f"{data.individual_name} (Passport No.: {data.passport_number})" if data.passport_number else data.individual_name
    enclosure_para = doc.add_paragraph(f'Enclosure: Statement of Net Worth of Mr./Ms. {name_with_passport} as at {data.certificate_date}')
    enclosure_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add page break for annexures
    doc.add_page_break()
    
    # Generate annexures
    generate_annexures(doc, data)
    
    return doc

def generate_annexures(doc, data: NetWorthData):
    """Generate all annexures"""
    
    # Annexure Header
    title = doc.add_paragraph()
    title_run = title.add_run('Annexure – Statement of Net Worth')
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()

    name_with_passport = f"{data.individual_name} (Passport No.: {data.passport_number})" if data.passport_number else data.individual_name
    name_para = doc.add_paragraph(f'Name of Individual: {name_with_passport}')
    name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    address_para = doc.add_paragraph(f'Address: {data.individual_address}')
    address_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    date_para = doc.add_paragraph(f'Date of Certificate: {data.certificate_date}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    purpose_para = doc.add_paragraph(f'Purpose: VISA Application – Submission to {data.embassy_name}')
    purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run(f'Statement of Net Worth as at {data.certificate_date}')
    run.bold = True
    
    doc.add_paragraph()
    
    # Annexure (i) - Movable Assets Summary
    para = doc.add_paragraph()
    run = para.add_run('Annexure (i) - Movable Assets')
    run.bold = True
    
    # Build list of categories with data
    categories = []
    if data.bank_accounts:
        categories.append(('1', 'Bank Account', 'A', data.total_bank_balance_inr, data.total_bank_balance_foreign))
    if data.insurance_policies:
        categories.append(('2', 'LIC', 'B', data.total_insurance_inr, data.total_insurance_foreign))
    if data.pf_accounts:
        categories.append(('3', 'P.F. Account', 'C', data.total_pf_accounts_inr, data.total_pf_accounts_foreign))
    if data.deposits:
        categories.append(('4', 'Deposit', 'D', data.total_deposits_inr, data.total_deposits_foreign))
    if data.nps_accounts:
        categories.append(('5', 'NPS', 'E', data.total_nps_inr, data.total_nps_foreign))
    if data.mutual_funds:
        categories.append(('6', 'Investment in Mutual Fund', 'F', data.total_mutual_funds_inr, data.total_mutual_funds_foreign))
    if data.shares:
        categories.append(('7', 'Shares & Securities', 'G', data.total_shares_inr, data.total_shares_foreign))
    if data.vehicles:
        categories.append(('8', 'Vehicles', 'H', data.total_vehicles_inr, data.total_vehicles_foreign))
    if data.post_office_schemes:
        categories.append(('9', 'Post Office Schemes', 'I', data.total_post_office_inr, data.total_post_office_foreign))
    if data.partnership_firms:
        categories.append(('10', 'Investments in Partnership Firms', 'J', data.total_partnership_firms_inr, data.total_partnership_firms_foreign))
    if data.gold_holdings:
        categories.append(('11', 'Gold', 'K', data.total_gold_inr, data.total_gold_foreign))
    
    # Only create table if there are categories with data
    if categories:
        # Create table with header + data rows + total row
        table = add_table_with_borders(doc, len(categories) + 2, 5)
    
        # Header row
        headers = ['Sr. No.', 'Particulars', 'Sub-Annexure', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
    
        # Data rows - only for categories with data
        for idx, (sr_no, particular, sub_annexure, inr_amount, foreign_amount) in enumerate(categories, 1):
            table.rows[idx].cells[0].text = sr_no
            table.rows[idx].cells[1].text = particular
            table.rows[idx].cells[2].text = sub_annexure
            table.rows[idx].cells[3].text = f'{inr_amount:,.2f}'
            table.rows[idx].cells[4].text = f'{foreign_amount:,.2f}'
        
        # Total row
        total_row = len(categories) + 1
        table.rows[total_row].cells[0].text = ''
        cell = table.rows[total_row].cells[1]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        table.rows[total_row].cells[2].text = ''
        cell = table.rows[total_row].cells[3]
        cell.text = f'{data.total_movable_assets_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[total_row].cells[4]
        cell.text = f'{data.total_movable_assets_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    else:
        # If no movable assets, add a note
        doc.add_paragraph('No movable assets to report.')
    
    # Sub Annexure A - Bank Accounts
    if data.bank_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (A) – Bank Account')
        run.bold = True
        table = add_table_with_borders(doc, len(data.bank_accounts) + 2, 6)
        
        # Headers
        headers = ['Sr. No.', 'Name of the Account Holder', 'Account No.', 'Bank Name', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Data rows
        for idx, acc in enumerate(data.bank_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = acc.holder_name
            table.rows[idx].cells[2].text = acc.account_number
            table.rows[idx].cells[3].text = acc.bank_name
            table.rows[idx].cells[4].text = f'{acc.balance_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{acc.balance_foreign:,.2f}'
        
        # Total row
        last_row = len(data.bank_accounts) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_bank_balance_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_bank_balance_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
    # Sub Annexure B - Insurance Policies
    if data.insurance_policies:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (B): LIFE INSURANCE POLICIES')
        run.bold = True
        table = add_table_with_borders(doc, len(data.insurance_policies) + 2, 5)
        
        headers = ['Sr. No.', 'POLICY Holder', 'Policy No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, policy in enumerate(data.insurance_policies, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = policy.holder_name
            table.rows[idx].cells[2].text = policy.policy_number
            table.rows[idx].cells[3].text = f'{policy.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{policy.amount_foreign:,.2f}'
        
        last_row = len(data.insurance_policies) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'TOTAL'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_insurance_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_insurance_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
    # Sub Annexure C - P.F. Accounts
    if data.pf_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (C) - P.F Account')
        run.bold = True
        table = add_table_with_borders(doc, len(data.pf_accounts) + 2, 5)
        headers = ['Sr. No.', 'Name of the Account Holder', 'PF Account No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, acc in enumerate(data.pf_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = acc.holder_name
            table.rows[idx].cells[2].text = acc.pf_account_number
            table.rows[idx].cells[3].text = f'{acc.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{acc.amount_foreign:,.2f}'
        
        last_row = len(data.pf_accounts) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_pf_accounts_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_pf_accounts_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True

    # Sub Annexure D - Deposits
    if data.deposits:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (D): Deposit')
        run.bold = True
        table = add_table_with_borders(doc, len(data.deposits) + 2, 5)
        headers = ['Sr. No.', 'Name of Investment Holder', 'A/C Number', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, dep in enumerate(data.deposits, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = dep.holder_name
            table.rows[idx].cells[2].text = dep.account_number
            table.rows[idx].cells[3].text = f'{dep.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{dep.amount_foreign:,.2f}'
        
        last_row = len(data.deposits) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_deposits_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_deposits_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
    # Sub Annexure E - NPS
    if data.nps_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (E) - NPS')
        run.bold = True
        table = add_table_with_borders(doc, len(data.nps_accounts) + 2, 5)
        headers = ['Sr. No.', 'Name of Owner', 'PRAN No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, nps in enumerate(data.nps_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = nps.owner_name
            table.rows[idx].cells[2].text = nps.pran_number
            table.rows[idx].cells[3].text = f'{nps.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{nps.amount_foreign:,.2f}'
        
        last_row = len(data.nps_accounts) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_nps_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_nps_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True

    # Sub Annexure F - Investment in Mutual Fund
    if data.mutual_funds:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (F) - Investment in Mutual Fund')
        run.bold = True
        table = add_table_with_borders(doc, len(data.mutual_funds) + 2, 6)
        headers = ['Sr. No.', 'Name of the Account Holder', 'Policy/Folio Number', 'Policy Name', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        for idx, mf in enumerate(data.mutual_funds, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = mf.holder_name
            table.rows[idx].cells[2].text = mf.folio_number
            table.rows[idx].cells[3].text = mf.policy_name
            table.rows[idx].cells[4].text = f'{mf.amount_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{mf.amount_foreign:,.2f}'

        last_row = len(data.mutual_funds) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_mutual_funds_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_mutual_funds_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True

    # Sub Annexure G - Gold
    if data.gold_holdings:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (G) - Gold')
        run.bold = True
        table = add_table_with_borders(doc, len(data.gold_holdings) + 2, 6)
        
        headers = ['Sr. No.', 'Name of Party', 'Weight (gram)', 'Rate/10 g (Rs.)', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, gold in enumerate(data.gold_holdings, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = gold.owner_name
            table.rows[idx].cells[2].text = f'{gold.weight_grams:.3f}'
            table.rows[idx].cells[3].text = f'{gold.rate_per_10g:,.2f}'
            table.rows[idx].cells[4].text = f'{gold.amount_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{gold.amount_foreign:,.2f}'
        
        last_row = len(data.gold_holdings) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_gold_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_gold_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        if data.gold_holdings and data.gold_holdings[0].valuation_date:
            valuer_para = doc.add_paragraph(f'As per the Property Valuation Certificates dated {data.gold_holdings[0].valuation_date} issued by Approved Valuer {data.gold_holdings[0].valuer_name}')
            valuer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Annexure (ii) - Immovable Assets
    if data.properties:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Annexure (ii) - Immovable Assets')
        run.bold = True
        table = add_table_with_borders(doc, len(data.properties) + 2, 4)
        
        headers = ['Sr. No.', 'Particulars of Property', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, prop in enumerate(data.properties, 1):
            table.rows[idx].cells[0].text = str(idx)
            property_text = f'{prop.owner_name}\n\n{prop.property_type}\n\n{prop.address}'
            if prop.valuation_date and prop.valuer_name:
                property_text += f'\n\n(Valuation as on {prop.valuation_date} by {prop.valuer_name})'
            table.rows[idx].cells[1].text = property_text
            table.rows[idx].cells[2].text = f'{prop.valuation_inr:,.2f}'
            table.rows[idx].cells[3].text = f'{prop.valuation_foreign:,.2f}'
        
        last_row = len(data.properties) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[2]
        cell.text = f'{data.total_immovable_assets_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_immovable_assets_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
    # Annexure (iii) - Liabilities
    if data.liabilities:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Annexure (iii) - Liabilities')
        run.bold = True
        table = add_table_with_borders(doc, len(data.liabilities) + 2, 4)
        
        headers = ['Sr. No.', 'Description', 'Details', 'Amount in INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for idx, liab in enumerate(data.liabilities, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = liab.description
            table.rows[idx].cells[2].text = liab.details
            table.rows[idx].cells[3].text = f'{liab.amount_inr:,.2f}'
        
        last_row = len(data.liabilities) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_liabilities_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
    # Net Worth Calculation
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f'Net Worth: ₹{data.net_worth_inr:,.2f}')
    run.bold = True
    run.font.size = Pt(14)
    
    # Notes
    doc.add_paragraph()

    notes_title = doc.add_paragraph('Notes:')
    notes_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note1 = doc.add_paragraph('1. The above Statement is prepared based on details and supporting documents provided by the individual.')
    note1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note2 = doc.add_paragraph('2. Valuation of assets is based on self-declaration / available records and has not been independently verified unless specified.')
    note2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note3 = doc.add_paragraph(f'3. This Annexure should be read with the Certificate dated {data.certificate_date} issued by the undersigned.')
    note3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Final Signature
    final_firm_para = doc.add_paragraph(f'FOR, {data.ca_firm_name.upper()}')
    final_firm_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    final_ca_para = doc.add_paragraph('CHARTERED ACCOUNTANTS')
    final_ca_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    final_frn_para = doc.add_paragraph(f'FRN: {data.ca_frn}')
    final_frn_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    # Create two-column layout for signature
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = True

    # Make table borders invisible using XML
    for row in sig_table.rows:
        for cell in row.cells:
            # Set all borders to nil (invisible)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                  r'<w:top w:val="nil"/>'
                                  r'<w:left w:val="nil"/>'
                                  r'<w:bottom w:val="nil"/>'
                                  r'<w:right w:val="nil"/>'
                                  r'</w:tcBorders>')
            tcPr.append(tcBorders)

    # Left column - CA details
    sig_table.rows[0].cells[0].text = f'{data.ca_partner_name.upper()}'
    sig_table.rows[1].cells[0].text = f'{data.ca_designation.upper()}'
    sig_table.rows[2].cells[0].text = f'MEMBERSHIP NO.: {data.ca_membership_no}'
    sig_table.rows[3].cells[0].text = 'UDIN: [TO BE GENERATED]'

    # Right column - Date and Place
    sig_table.rows[0].cells[1].text = f'DATE: {data.certificate_date}'
    sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sig_table.rows[1].cells[1].text = f'PLACE: {data.ca_place.upper()}'
    sig_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Leave bottom two cells in right column empty for spacing
    sig_table.rows[2].cells[1].text = ''
    sig_table.rows[3].cells[1].text = ''

def convert_to_words(number):
    """Convert number to words (simplified - for production use num2words library)"""
    # This is a simplified version - in production, use num2words library
    return f"{number:,.2f}"

# ==================== STREAMLIT UI ====================

def auto_fill_test_data():
    """Auto-fill all form fields with comprehensive test data for testing"""
    import datetime

    # Initialize session state data
    st.session_state.data = NetWorthData(
        individual_name="Bharatkumar Dhulabhai Patel",
        individual_address="29/B, Ratnamani Tenaments, Ahmedabad, Gujarat - 380001",
        certificate_date=datetime.date.today().strftime("%d/%m/%Y"),
        engagement_date=datetime.date.today().strftime("%d/%m/%Y"),
        embassy_name="Canadian High Commission",
        embassy_address="7/8 Shantipath, Chanakyapuri\nNew Delhi - 110021",
        passport_number="A12345678",
        foreign_currency="CAD",
        exchange_rate=63.34
    )

    # Bank Accounts
    st.session_state.data.bank_accounts = [
        BankAccount(
            holder_name="Bharatkumar Dhulabhai Patel",
            account_number="10733415306",
            bank_name="State Bank of India",
            balance_inr=1078118.46,
            statement_date=datetime.date.today().strftime("%d/%m/%Y")
        ),
        BankAccount(
            holder_name="Bharatkumar Dhulabhai Patel",
            account_number="20093766850",
            bank_name="HDFC Bank",
            balance_inr=385989.69,
            statement_date=datetime.date.today().strftime("%d/%m/%Y")
        )
    ]

    # Insurance Policies
    st.session_state.data.insurance_policies = [
        InsurancePolicy(
            holder_name="Bharatkumar Dhulabhai Patel",
            policy_number="71234567890",
            amount_inr=253618.00
        )
    ]

    # P.F. Accounts
    st.session_state.data.pf_accounts = [
        PFAccount(
            holder_name="Bharatkumar Dhulabhai Patel",
            pf_account_number="GJ/AMD/12345/678",
            amount_inr=1850652.00
        )
    ]

    # Deposits
    st.session_state.data.deposits = [
        Deposit(
            holder_name="Bharatkumar Dhulabhai Patel",
            account_number="FD123456789",
            amount_inr=1000000.00
        )
    ]

    # NPS Accounts
    st.session_state.data.nps_accounts = [
        NPSAccount(
            owner_name="Bharatkumar Dhulabhai Patel",
            pran_number="110021379979",
            amount_inr=2578706.52
        )
    ]

    # Mutual Funds
    st.session_state.data.mutual_funds = [
        MutualFund(
            holder_name="Bharatkumar Dhulabhai Patel",
            folio_number="34521763/59",
            policy_name="HDFC Flexi Cap Fund",
            amount_inr=126662.88
        )
    ]

    # Shares
    st.session_state.data.shares = [
        Share(
            company_name="Reliance Industries Ltd",
            num_shares=100,
            market_price_inr=2450.00
        ),
        Share(
            company_name="TCS Ltd",
            num_shares=50,
            market_price_inr=3200.00
        )
    ]

    # Vehicles
    st.session_state.data.vehicles = [
        Vehicle(
            vehicle_type="Car",
            make_model_year="Toyota Innova Crysta 2020",
            registration_number="GJ01AB1234",
            market_value_inr=1500000.00
        )
    ]

    # Post Office Schemes
    st.session_state.data.post_office_schemes = [
        PostOfficeScheme(
            scheme_type="National Savings Certificate (NSC)",
            account_number="NSC123456789",
            amount_inr=500000.00
        )
    ]

    # Partnership Firms
    st.session_state.data.partnership_firms = [
        PartnershipFirm(
            firm_name="Patel Brothers Trading Co.",
            partner_name="Bharatkumar Dhulabhai Patel",
            holding_percentage=33.33,
            capital_balance_inr=2000000.00,
            valuation_date=datetime.date.today().strftime("%d/%m/%Y")
        )
    ]

    # Gold Holdings
    st.session_state.data.gold_holdings = [
        GoldHolding(
            owner_name="Bharatkumar Dhulabhai Patel",
            weight_grams=399.570,
            rate_per_10g=109500,
            valuation_date=datetime.date.today().strftime("%d/%m/%Y"),
            valuer_name="Approved Valuer - Ahmedabad"
        )
    ]

    # Properties
    st.session_state.data.properties = [
        Property(
            owner_name="Bharatkumar Dhulabhai Patel",
            property_type="Residential House",
            address="29/B, Ratnamani Tenaments, Survey No. 123, Ahmedabad",
            valuation_inr=1260000,
            valuation_date=datetime.date.today().strftime("%d/%m/%Y"),
            valuer_name="Approved Valuer - Ahmedabad"
        ),
        Property(
            owner_name="Bharatkumar Dhulabhai Patel",
            property_type="Commercial Property",
            address="Shop No. 5, Commercial Complex, SG Highway, Ahmedabad",
            valuation_inr=2500000,
            valuation_date=datetime.date.today().strftime("%d/%m/%Y"),
            valuer_name="Approved Valuer - Ahmedabad"
        )
    ]

    # Liabilities
    st.session_state.data.liabilities = [
        Liability(
            description="Home Loan",
            amount_inr=800000.00,
            details="Housing loan from SBI for residential property"
        )
    ]

    print("✅ Test data auto-filled successfully!")
    print(f"📊 Total Assets: ₹{st.session_state.data.total_movable_assets_inr + st.session_state.data.total_immovable_assets_inr:,.2f}")
    print(f"💰 Net Worth: ₹{st.session_state.data.net_worth_inr:,.2f}")
    print("🎯 Ready to generate certificate on Summary page!")

def main():
    st.set_page_config(page_title="Net Worth Certificate Generator", layout="wide", page_icon="📄")
    
    # Check for test mode parameter
    query_params = st.query_params
    test_mode = query_params.get('test', 'false').lower() == 'true'

    # Theme management
    if 'theme' not in st.session_state:
        st.session_state.theme = 'light'

    # Auto-fill test data if in test mode
    if test_mode and 'data' not in st.session_state:
        auto_fill_test_data()

        # Auto-navigate to summary page in test mode
        if test_mode:
            st.markdown("""
            <script>
                // Auto-navigate to Summary tab
                setTimeout(() => {
                    const tabs = document.querySelectorAll('[data-baseweb="tab"]');
                    if (tabs.length > 0) {
                        tabs[tabs.length - 1].click(); // Click last tab (Summary)
                    }
                }, 1000); // Wait 1 second for data to load
            </script>
            """, unsafe_allow_html=True)

    # Theme toggle function
    def toggle_theme():
        st.session_state.theme = 'dark' if st.session_state.theme == 'light' else 'light'
        st.rerun()

    # Add theme selector in sidebar
    with st.sidebar:
        st.markdown("### 🎨 Theme Settings")
        theme_options = ["☀️ Light Mode", "🌙 Dark Mode"]
        current_theme_display = "☀️ Light Mode" if st.session_state.theme == 'light' else "🌙 Dark Mode"

        if st.button(f"Switch to {theme_options[1] if st.session_state.theme == 'light' else theme_options[0]}",
                     key="theme_toggle", use_container_width=True):
            toggle_theme()

        # Add extra spacing after theme toggle
        st.markdown("<br><br>", unsafe_allow_html=True)

    # Add custom CSS for Quicksand font and themes
    light_theme_css = """
    @import url('https://fonts.googleapis.com/css2?family=Quicksand:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTitle {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 600 !important;
    }

    .stHeader {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stSubheader {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stText {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stMarkdown {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stButton button {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stTextInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stNumberInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTextArea textarea {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stSelectbox select {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stDateInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stMetric {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stExpander {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTabs {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTab {
        font-family: 'Quicksand', sans-serif !important;
    }

    /* Light theme styles */
    .main .block-container {
        background-color: #ffffff !important;
        color: #212529 !important;
    }

    .stApp {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%) !important;
    }

    .stTitle, .stHeader, .stSubheader {
        color: #212529 !important;
        font-weight: 600 !important;
    }

    .stText {
        color: #495057 !important;
    }

    .stMarkdown {
        color: #495057 !important;
    }

    .stMarkdown p {
        color: #495057 !important;
    }

    .stMarkdown strong, .stMarkdown b {
        color: #212529 !important;
        font-weight: 600 !important;
    }

    /* Ensure all header-like elements are dark */
    h1, h2, h3, h4, h5, h6 {
        color: #212529 !important;
    }

    /* Bold text in general */
    b, strong {
        color: #212529 !important;
        font-weight: 600 !important;
    }

    .stButton button {
        background-color: #007bff !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
    }

    .stButton button:hover {
        background-color: #0056b3 !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.15) !important;
    }

    .stTextInput input, .stNumberInput input, .stTextArea textarea, .stSelectbox select {
        background-color: #ffffff !important;
        color: #495057 !important;
        border: 2px solid #ced4da !important;
        border-radius: 6px !important;
        padding: 12px !important;
    }

    .stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus, .stSelectbox select:focus {
        border-color: #007bff !important;
        box-shadow: 0 0 0 0.2rem rgba(0, 123, 255, 0.25) !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        background-color: #ffffff !important;
        border-radius: 8px !important;
        border: 1px solid #dee2e6 !important;
        padding: 8px !important;
    }

    .stTabs [data-baseweb="tab"] {
        color: #495057 !important;
        background-color: transparent !important;
        border: none !important;
        padding: 8px 16px !important;
        margin: 2px !important;
        border-radius: 6px !important;
    }

    .stTabs [aria-selected="true"] {
        background-color: #007bff !important;
        color: white !important;
    }

    .stMetric {
        background-color: #ffffff !important;
        border: 1px solid #dee2e6 !important;
        border-radius: 8px !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05) !important;
    }

    .stMetric .metric-value {
        color: #007bff !important;
    }

    .stExpander {
        background-color: #ffffff !important;
        border: 1px solid #dee2e6 !important;
        border-radius: 8px !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05) !important;
    }

    .stExpander header {
        background-color: #f8f9fa !important;
        border-radius: 8px 8px 0 0 !important;
    }

    .stAlert {
        background-color: #d1ecf1 !important;
        color: #0c5460 !important;
        border: 1px solid #bee5eb !important;
        border-radius: 6px !important;
    }

    .sidebar .sidebar-content {
        background-color: #ffffff !important;
        border-right: 1px solid #dee2e6 !important;
    }

    .sidebar .sidebar-content .stMarkdown {
        color: #495057 !important;
    }

    /* Form labels and help text */
    .stTextInput label, .stNumberInput label, .stTextArea label, .stSelectbox label {
        color: #495057 !important;
        font-weight: 500 !important;
    }

    /* Success/Error messages */
    .stSuccess {
        background-color: #d4edda !important;
        color: #155724 !important;
        border: 1px solid #c3e6cb !important;
    }

    .stError {
        background-color: #f8d7da !important;
        color: #721c24 !important;
        border: 1px solid #f5c6cb !important;
    }

    /* Checkbox and radio styling */
    .stCheckbox, .stRadio {
        color: #495057 !important;
    }

    /* Progress bars */
    .stProgress > div > div {
        background-color: #007bff !important;
    }

    /* Additional text elements that might be white */
    .css-1v0mbdj, .css-1v0mbdj * {
        color: #495057 !important;
    }

    /* Ensure all text in containers is properly colored */
    .css-1d391kg, .css-1d391kg * {
        color: #495057 !important;
    }

    /* Fix for any remaining light text */
    [data-testid="stMarkdownContainer"] p {
        color: #495057 !important;
    }

    [data-testid="stMarkdownContainer"] strong {
        color: #212529 !important;
    }

    /* Tab content text */
    .stTabs [data-baseweb="tab-panel"] {
        color: #495057 !important;
    }

    /* Sidebar text */
    .sidebar .sidebar-content * {
        color: #495057 !important;
    }

    /* Form help text and descriptions */
    .stTextInput div small, .stNumberInput div small, .stTextArea div small {
        color: #6c757d !important;
    }
    """

    dark_theme_css = """
    @import url('https://fonts.googleapis.com/css2?family=Quicksand:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTitle {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 600 !important;
    }

    .stHeader {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stSubheader {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stText {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stMarkdown {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stButton button {
        font-family: 'Quicksand', sans-serif !important;
        font-weight: 500 !important;
    }

    .stTextInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stNumberInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTextArea textarea {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stSelectbox select {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stDateInput input {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stMetric {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stExpander {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTabs {
        font-family: 'Quicksand', sans-serif !important;
    }

    .stTab {
        font-family: 'Quicksand', sans-serif !important;
    }

    /* Dark theme styles */
    .main .block-container {
        background-color: #1a1a1a !important;
        color: #ffffff !important;
    }

    .stApp {
        background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%) !important;
    }

    .stTitle, .stHeader, .stSubheader {
        color: #ffffff !important;
    }

    .stText {
        color: #ffffff !important;
    }

    .stMarkdown {
        color: #ffffff !important;
    }

    .stButton button {
        background-color: #007bff !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 500 !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
    }

    .stButton button:hover {
        background-color: #0056b3 !important;
        box-shadow: 0 4px 8px rgba(0,0,0,0.15) !important;
    }

    .stTextInput input, .stNumberInput input, .stTextArea textarea, .stSelectbox select {
        background-color: #2d2d2d !important;
        color: #ffffff !important;
        border: 1px solid #404040 !important;
        border-radius: 6px !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        background-color: #2d2d2d !important;
        border-radius: 8px !important;
    }

    .stTabs [data-baseweb="tab"] {
        color: #ffffff !important;
    }

    .stTabs [aria-selected="true"] {
        background-color: #4CAF50 !important;
        color: white !important;
    }

    .stMetric {
        background-color: #2d2d2d !important;
        border: 1px solid #404040 !important;
        border-radius: 8px !important;
    }

    .stExpander {
        background-color: #2d2d2d !important;
        border: 1px solid #404040 !important;
        border-radius: 8px !important;
    }

    .stAlert {
        background-color: #1e3a8a !important;
        color: #93c5fd !important;
        border: 1px solid #3b82f6 !important;
    }

    .sidebar .sidebar-content {
        background-color: #2d2d2d !important;
    }
    """

    # Apply the appropriate theme CSS
    if st.session_state.theme == 'light':
        st.markdown(f"<style>{light_theme_css}</style>", unsafe_allow_html=True)
    else:
        st.markdown(f"<style>{dark_theme_css}</style>", unsafe_allow_html=True)

    # Enhanced branding with optiCERTIFICATE logo
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image("Logo.png", use_column_width='auto')

    st.markdown("---")
    
    # Initialize session state
    if 'data' not in st.session_state:
        st.session_state.data = NetWorthData(
            individual_name="",
            individual_address="",
            certificate_date=datetime.date.today().strftime("%d/%m/%Y"),
            engagement_date=datetime.date.today().strftime("%d/%m/%Y"),
            embassy_name="",
            embassy_address="",
            passport_number=""
        )
    
    # For forward compatibility: ensure new fields exist on older session state objects
    new_fields = {
        'pf_accounts': [], 'deposits': [], 'nps_accounts': [], 'mutual_funds': [],
        'shares': [], 'vehicles': [], 'post_office_schemes': [], 'partnership_firms': [],
        'passport_number': ''
    }
    for field, default_value in new_fields.items():
        if not hasattr(st.session_state.data, field):
            setattr(st.session_state.data, field, default_value)
    
    # Tabs for different sections
    tabs = st.tabs([
        "📋 Basic Info", 
        "🏦 Bank Accounts", 
        "🛡️ Insurance", 
        "📈 P.F. Accounts",
        "💰 Deposits",
        "📊 NPS",
        "💼 Mutual Funds",
        "📈 Shares",
        "🚗 Vehicles",
        "📮 Post Office",
        "🤝 Partnership Firms",
        "💎 Gold/Valuables", 
        "🏠 Properties", 
        "💳 Liabilities",
        "📊 Summary & Generate"
    ])
    
    # Tab 1: Basic Information
    with tabs[0]:
        st.header("Basic Information")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Personal Details")
            st.session_state.data.individual_name = st.text_input(
                "Individual's Full Name *", 
                value=st.session_state.data.individual_name,
                key="ind_name"
            )
            st.session_state.data.passport_number = st.text_input(
                "Passport Number", 
                value=st.session_state.data.passport_number,
                key="passport_num",
                help="Passport number will be displayed after the name in the certificate"
            )
            st.session_state.data.individual_address = st.text_area(
                "Individual's Address *", 
                value=st.session_state.data.individual_address,
                height=100,
                key="ind_addr"
            )
            
            st.subheader("Certificate Dates")
            cert_date = st.date_input("Certificate Date", value=datetime.date.today())
            st.session_state.data.certificate_date = cert_date.strftime("%d/%m/%Y")
            
            eng_date = st.date_input("Engagement Letter Date", value=datetime.date.today())
            st.session_state.data.engagement_date = eng_date.strftime("%d/%m/%Y")
        
        with col2:
            st.subheader("Embassy/Consulate Details")
            st.session_state.data.embassy_name = st.text_input(
                "Embassy/Consulate Name *", 
                value=st.session_state.data.embassy_name,
                key="emb_name"
            )
            st.session_state.data.embassy_address = st.text_area(
                "Embassy Address *", 
                value=st.session_state.data.embassy_address,
                height=100,
                key="emb_addr"
            )
            
            st.subheader("Currency Settings")
            st.session_state.data.foreign_currency = st.selectbox(
                "Foreign Currency",
                ["CAD", "USD", "EUR", "GBP", "AUD"],
                index=0
            )
            st.session_state.data.exchange_rate = st.number_input(
                f"Exchange Rate (INR to 1 {st.session_state.data.foreign_currency})",
                min_value=0.01,
                value=63.34,
                step=0.01,
                format="%.2f"
            )
            st.session_state.exchange_rate = st.session_state.data.exchange_rate
        
        st.markdown("---")
        st.subheader("Chartered Accountant Details")
        
        col1_ca, col2_ca = st.columns(2)
        with col1_ca:
            selected_ca_name = st.selectbox(
                "Select Signing Partner",
                options=list(CA_PARTNERS.keys()),
                key="ca_name_select"
            )
            # Update session state with selected CA details
            st.session_state.data.ca_partner_name = selected_ca_name
            st.session_state.data.ca_membership_no = CA_PARTNERS[selected_ca_name]["membership_no"]

        with col2_ca:
            st.info(f"""
            **Firm:** {st.session_state.data.ca_firm_name}
            **FRN:** {st.session_state.data.ca_frn}
            """)
    
    # Tab 2: Bank Accounts
    with tabs[1]:
        st.header("Bank Accounts")
        
        st.subheader("Add New Bank Account")
        with st.form("bank_account_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                holder_name = st.text_input("Account Holder Name")
                account_number = st.text_input("Account Number")
            with col2:
                bank_name = st.text_input("Bank Name")
                balance = st.number_input("Balance (INR)", min_value=0.0, step=1000.0, format="%.2f")
            
            statement_date = st.date_input("Statement Date", value=datetime.date.today())
            
            if st.form_submit_button("➕ Add Bank Account"):
                if holder_name and account_number and bank_name:
                    new_account = BankAccount(
                        holder_name=holder_name,
                        account_number=account_number,
                        bank_name=bank_name,
                        balance_inr=balance,
                        statement_date=statement_date.strftime("%d/%m/%Y")
                    )
                    st.session_state.data.bank_accounts.append(new_account)
                    st.success("✅ Bank account added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Bank Accounts")
        if st.session_state.data.bank_accounts:
            for idx, acc in enumerate(st.session_state.data.bank_accounts):
                with st.expander(f"Account {idx+1}: {acc.bank_name} - {acc.account_number}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Holder:** {acc.holder_name}")
                        st.write(f"**Bank:** {acc.bank_name}")
                    with col2:
                        st.write(f"**Balance (INR):** ₹{acc.balance_inr:,.2f}")
                        st.write(f"**Balance ({st.session_state.data.foreign_currency}):** {acc.balance_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_bank_{idx}"):
                            st.session_state.data.bank_accounts.pop(idx)
                            st.rerun()
            
            st.metric("Total Bank Balance (INR)", f"₹{st.session_state.data.total_bank_balance_inr:,.2f}")
        else:
            st.info("No bank accounts added yet")
    
    # Tab 3: Insurance Policies
    with tabs[2]:
        st.header("Life Insurance Policies")
        
        st.subheader("Add New Insurance Policy")
        with st.form("insurance_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                holder_name = st.text_input("Policy Holder Name")
            with col2:
                policy_number = st.text_input("Policy Number")
            with col3:
                amount = st.number_input("Surrender/Maturity Value (INR)", min_value=0.0, step=1000.0, format="%.2f")
            
            if st.form_submit_button("➕ Add Insurance Policy"):
                if holder_name and policy_number:
                    new_policy = InsurancePolicy(
                        holder_name=holder_name,
                        policy_number=policy_number,
                        amount_inr=amount
                    )
                    st.session_state.data.insurance_policies.append(new_policy)
                    st.success("✅ Insurance policy added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Insurance Policies")
        if st.session_state.data.insurance_policies:
            for idx, policy in enumerate(st.session_state.data.insurance_policies):
                with st.expander(f"Policy {idx+1}: {policy.policy_number}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Holder:** {policy.holder_name}")
                        st.write(f"**Policy No:** {policy.policy_number}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{policy.amount_inr:,.2f}")
                        st.write(f"**Amount ({st.session_state.data.foreign_currency}):** {policy.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_ins_{idx}"):
                            st.session_state.data.insurance_policies.pop(idx)
                            st.rerun()
            
            st.metric("Total Insurance Value (INR)", f"₹{st.session_state.data.total_insurance_inr:,.2f}")
        else:
            st.info("No insurance policies added yet")
    
    # Tab 4: P.F. Accounts
    with tabs[3]:
        st.header("Provident Fund (P.F.) Accounts")
        
        st.subheader("Add New P.F. Account")
        with st.form("pf_account_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                holder_name = st.text_input("Account Holder Name")
            with col2:
                pf_number = st.text_input("P.F. Account Number")
            with col3:
                amount = st.number_input("Balance (INR)", min_value=0.0, step=1000.0, format="%.2f")
            
            if st.form_submit_button("➕ Add P.F. Account"):
                if holder_name and pf_number:
                    new_pf = PFAccount(
                        holder_name=holder_name,
                        pf_account_number=pf_number,
                        amount_inr=amount
                    )
                    st.session_state.data.pf_accounts.append(new_pf)
                    st.success("✅ P.F. account added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current P.F. Accounts")
        if st.session_state.data.pf_accounts:
            for idx, acc in enumerate(st.session_state.data.pf_accounts):
                with st.expander(f"P.F. Account {idx+1}: {acc.pf_account_number}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Holder:** {acc.holder_name}")
                    with col2:
                        st.write(f"**Balance (INR):** ₹{acc.amount_inr:,.2f}")
                        st.write(f"**Balance ({st.session_state.data.foreign_currency}):** {acc.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_pf_{idx}"):
                            st.session_state.data.pf_accounts.pop(idx)
                            st.rerun()
            
            st.metric("Total P.F. Balance (INR)", f"₹{st.session_state.data.total_pf_accounts_inr:,.2f}")
        else:
            st.info("No P.F. accounts added yet")

    # Tab 5: Deposits
    with tabs[4]:
        st.header("Deposits (FD, RD, etc.)")
        
        st.subheader("Add New Deposit")
        with st.form("deposit_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                holder_name = st.text_input("Investment Holder Name")
            with col2:
                acc_number = st.text_input("A/C Number")
            with col3:
                amount = st.number_input("Amount (INR)", min_value=0.0, step=1000.0, format="%.2f")
            
            if st.form_submit_button("➕ Add Deposit"):
                if holder_name and acc_number:
                    new_deposit = Deposit(
                        holder_name=holder_name,
                        account_number=acc_number,
                        amount_inr=amount
                    )
                    st.session_state.data.deposits.append(new_deposit)
                    st.success("✅ Deposit added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Deposits")
        if st.session_state.data.deposits:
            for idx, dep in enumerate(st.session_state.data.deposits):
                with st.expander(f"Deposit {idx+1}: {dep.account_number}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Holder:** {dep.holder_name}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{dep.amount_inr:,.2f}")
                        st.write(f"**Amount ({st.session_state.data.foreign_currency}):** {dep.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_dep_{idx}"):
                            st.session_state.data.deposits.pop(idx)
                            st.rerun()
            
            st.metric("Total Deposit Value (INR)", f"₹{st.session_state.data.total_deposits_inr:,.2f}")
        else:
            st.info("No deposits added yet")

    # Tab 6: NPS
    with tabs[5]:
        st.header("National Pension System (NPS)")
        
        st.subheader("Add New NPS Account")
        with st.form("nps_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                owner_name = st.text_input("Name of Owner")
            with col2:
                pran_number = st.text_input("PRAN No.")
            with col3:
                amount = st.number_input("Amount (INR)", min_value=0.0, step=1000.0, format="%.2f")
            
            if st.form_submit_button("➕ Add NPS Account"):
                if owner_name and pran_number:
                    new_nps = NPSAccount(
                        owner_name=owner_name,
                        pran_number=pran_number,
                        amount_inr=amount
                    )
                    st.session_state.data.nps_accounts.append(new_nps)
                    st.success("✅ NPS account added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current NPS Accounts")
        if st.session_state.data.nps_accounts:
            for idx, nps in enumerate(st.session_state.data.nps_accounts):
                with st.expander(f"NPS Account {idx+1}: {nps.pran_number}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Owner:** {nps.owner_name}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{nps.amount_inr:,.2f}")
                        st.write(f"**Amount ({st.session_state.data.foreign_currency}):** {nps.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_nps_{idx}"):
                            st.session_state.data.nps_accounts.pop(idx)
                            st.rerun()
            
            st.metric("Total NPS Value (INR)", f"₹{st.session_state.data.total_nps_inr:,.2f}")
        else:
            st.info("No NPS accounts added yet")

    # Tab 7: Mutual Funds
    with tabs[6]:
        st.header("Investment in Mutual Funds")

        st.subheader("Add New Mutual Fund")
        with st.form("mf_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                holder_name = st.text_input("Name of the Account Holder")
                folio_number = st.text_input("Policy/Folio Number")
            with col2:
                policy_name = st.text_input("Policy Name")
                amount = st.number_input("Amount (INR)", min_value=0.0, step=1000.0, format="%.2f")

            if st.form_submit_button("➕ Add Mutual Fund"):
                if holder_name and folio_number and policy_name:
                    new_mf = MutualFund(
                        holder_name=holder_name,
                        folio_number=folio_number,
                        policy_name=policy_name,
                        amount_inr=amount
                    )
                    st.session_state.data.mutual_funds.append(new_mf)
                    st.success("✅ Mutual fund added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")

        st.subheader("Current Mutual Funds")
        if st.session_state.data.mutual_funds:
            for idx, mf in enumerate(st.session_state.data.mutual_funds):
                with st.expander(f"Mutual Fund {idx+1}: {mf.policy_name}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Holder:** {mf.holder_name}")
                        st.write(f"**Folio No:** {mf.folio_number}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{mf.amount_inr:,.2f}")
                        st.write(f"**Amount ({st.session_state.data.foreign_currency}):** {mf.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_mf_{idx}"):
                            st.session_state.data.mutual_funds.pop(idx)
                            st.rerun()
            
            st.metric("Total Mutual Fund Value (INR)", f"₹{st.session_state.data.total_mutual_funds_inr:,.2f}")
        else:
            st.info("No mutual funds added yet")

    # Tab 8: Shares
    with tabs[7]:
        st.header("Shares & Securities")
        st.subheader("Add New Share Holding")
        with st.form("share_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                company_name = st.text_input("Company Name")
            with col2:
                num_shares = st.number_input("Number of Shares", min_value=1, step=1)
            with col3:
                price = st.number_input("Market Price per Share (INR)", min_value=0.0, format="%.2f")
            
            if st.form_submit_button("➕ Add Share Holding"):
                if company_name and num_shares > 0:
                    new_share = Share(company_name=company_name, num_shares=num_shares, market_price_inr=price)
                    st.session_state.data.shares.append(new_share)
                    st.success("✅ Share holding added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Share Holdings")
        if st.session_state.data.shares:
            for idx, s in enumerate(st.session_state.data.shares):
                with st.expander(f"Share {idx+1}: {s.company_name}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Company:** {s.company_name}")
                        st.write(f"**Shares:** {s.num_shares}")
                        st.write(f"**Price/Share (INR):** ₹{s.market_price_inr:,.2f}")
                    with col2:
                        st.write(f"**Total Value (INR):** ₹{s.amount_inr:,.2f}")
                        st.write(f"**Total Value ({st.session_state.data.foreign_currency}):** {s.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_share_{idx}"):
                            st.session_state.data.shares.pop(idx)
                            st.rerun()
            
            st.metric("Total Shares Value (INR)", f"₹{st.session_state.data.total_shares_inr:,.2f}")
        else:
            st.info("No shares added yet")

    # Tab 9: Vehicles
    with tabs[8]:
        st.header("Vehicles")
        st.subheader("Add New Vehicle")
        with st.form("vehicle_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                vehicle_type = st.selectbox("Vehicle Type", ["Car", "Motorcycle", "Scooter"])
                make_model_year = st.text_input("Make, Model & Year")
            with col2:
                reg_number = st.text_input("Registration Number")
                market_value = st.number_input("Estimated Market Value (INR)", min_value=0.0, format="%.2f")

            if st.form_submit_button("➕ Add Vehicle"):
                if vehicle_type and make_model_year and reg_number:
                    new_vehicle = Vehicle(vehicle_type=vehicle_type, make_model_year=make_model_year, registration_number=reg_number, market_value_inr=market_value)
                    st.session_state.data.vehicles.append(new_vehicle)
                    st.success("✅ Vehicle added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Vehicles")
        if st.session_state.data.vehicles:
            for idx, v in enumerate(st.session_state.data.vehicles):
                with st.expander(f"Vehicle {idx+1}: {v.vehicle_type} - {v.make_model_year}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Type:** {v.vehicle_type}")
                        st.write(f"**Make/Model/Year:** {v.make_model_year}")
                        st.write(f"**Reg No:** {v.registration_number}")
                    with col2:
                        st.write(f"**Market Value (INR):** ₹{v.market_value_inr:,.2f}")
                        st.write(f"**Market Value ({st.session_state.data.foreign_currency}):** {v.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_vehicle_{idx}"):
                            st.session_state.data.vehicles.pop(idx)
                            st.rerun()
            
            st.metric("Total Vehicle Value (INR)", f"₹{st.session_state.data.total_vehicles_inr:,.2f}")
        else:
            st.info("No vehicles added yet")

    # Tab 10: Post Office Schemes
    with tabs[9]:
        st.header("Post Office Schemes")
        st.subheader("Add New Post Office Scheme")
        with st.form("po_form", clear_on_submit=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                scheme_type = st.selectbox("Scheme Type", ["National Savings Certificate (NSC)", "Kisan Vikas Patra (KVP)", "Time Deposit", "Other"])
            with col2:
                acc_number = st.text_input("Certificate/Account Number")
            with col3:
                amount = st.number_input("Investment Amount (INR)", min_value=0.0, format="%.2f")

            if st.form_submit_button("➕ Add Scheme"):
                if scheme_type and acc_number:
                    new_scheme = PostOfficeScheme(scheme_type=scheme_type, account_number=acc_number, amount_inr=amount)
                    st.session_state.data.post_office_schemes.append(new_scheme)
                    st.success("✅ Scheme added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")

        st.subheader("Current Schemes")
        if st.session_state.data.post_office_schemes:
            for idx, p in enumerate(st.session_state.data.post_office_schemes):
                with st.expander(f"Scheme {idx+1}: {p.scheme_type}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Type:** {p.scheme_type}")
                        st.write(f"**Account No:** {p.account_number}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{p.amount_inr:,.2f}")
                        st.write(f"**Amount ({st.session_state.data.foreign_currency}):** {p.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_po_{idx}"):
                            st.session_state.data.post_office_schemes.pop(idx)
                            st.rerun()
            
            st.metric("Total Post Office Scheme Value (INR)", f"₹{st.session_state.data.total_post_office_inr:,.2f}")
        else:
            st.info("No schemes added yet")

    # Tab 11: Partnership Firms
    with tabs[10]:
        st.header("Investments in Partnership Firms")
        st.subheader("Add New Partnership Firm Investment")
        with st.form("firm_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                firm_name = st.text_input("Name of the Partnership Firm")
                partner_name = st.text_input("Name of the Partner")
                holding_percentage = st.number_input("Percentage of Holdings (%)", min_value=0.0, max_value=100.0, format="%.2f")
            with col2:
                capital_balance = st.number_input("Capital Account Balance (INR)", min_value=0.0, format="%.2f")
                val_date = st.date_input("Valuation Date", value=datetime.date.today())

            if st.form_submit_button("➕ Add Firm Investment"):
                if firm_name and partner_name:
                    new_firm = PartnershipFirm(
                        firm_name=firm_name,
                        partner_name=partner_name,
                        holding_percentage=holding_percentage,
                        capital_balance_inr=capital_balance,
                        valuation_date=val_date.strftime("%d/%m/%Y")
                    )
                    st.session_state.data.partnership_firms.append(new_firm)
                    st.success("✅ Firm investment added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Firm Investments")
        if st.session_state.data.partnership_firms:
            # UI to display firms
            st.metric("Total Partnership Firm Investment (INR)", f"₹{st.session_state.data.total_partnership_firms_inr:,.2f}")
        else:
            st.info("No firm investments added yet")

    # Tab 12: Gold/Valuables
    with tabs[11]:
        st.header("Gold & Valuables")
        
        st.subheader("Add Gold Holding")
        with st.form("gold_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                owner_name = st.text_input("Owner Name")
                weight = st.number_input("Weight (grams)", min_value=0.0, step=0.001, format="%.3f")
                rate = st.number_input("Rate per 10g (INR)", min_value=0.0, step=100.0, format="%.2f")
            with col2:
                val_date = st.date_input("Valuation Date", value=datetime.date.today())
                valuer_name = st.text_input("Valuer Name & Details")
            
            if st.form_submit_button("➕ Add Gold Holding"):
                if owner_name and weight > 0:
                    new_gold = GoldHolding(
                        owner_name=owner_name,
                        weight_grams=weight,
                        rate_per_10g=rate,
                        valuation_date=val_date.strftime("%d/%m/%Y"),
                        valuer_name=valuer_name
                    )
                    st.session_state.data.gold_holdings.append(new_gold)
                    st.success("✅ Gold holding added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Gold Holdings")
        if st.session_state.data.gold_holdings:
            for idx, gold in enumerate(st.session_state.data.gold_holdings):
                with st.expander(f"Gold {idx+1}: {gold.weight_grams}g"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Owner:** {gold.owner_name}")
                        st.write(f"**Weight:** {gold.weight_grams:.3f} grams")
                        st.write(f"**Rate:** ₹{gold.rate_per_10g:,.2f} per 10g")
                    with col2:
                        st.write(f"**Value (INR):** ₹{gold.amount_inr:,.2f}")
                        st.write(f"**Value ({st.session_state.data.foreign_currency}):** {gold.amount_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_gold_{idx}"):
                            st.session_state.data.gold_holdings.pop(idx)
                            st.rerun()
            
            st.metric("Total Gold Value (INR)", f"₹{st.session_state.data.total_gold_inr:,.2f}")
        else:
            st.info("No gold holdings added yet")
    
    # Tab 13: Properties
    with tabs[12]:
        st.header("Immovable Properties")
        
        st.subheader("Add Property")
        with st.form("property_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                owner_name = st.text_input("Owner Name")
                property_type = st.selectbox("Property Type", ["Residential House", "Commercial Property", "Agriculture Land", "Plot"])
                address = st.text_area("Complete Address (with Survey/Khata No.)", height=100)
            with col2:
                valuation = st.number_input("Valuation (INR)", min_value=0.0, step=10000.0, format="%.2f")
                val_date = st.date_input("Valuation Date", value=datetime.date.today())
                valuer_name = st.text_input("Valuer Name & Details")
            
            if st.form_submit_button("➕ Add Property"):
                if owner_name and address:
                    new_property = Property(
                        owner_name=owner_name,
                        property_type=property_type,
                        address=address,
                        valuation_inr=valuation,
                        valuation_date=val_date.strftime("%d/%m/%Y"),
                        valuer_name=valuer_name
                    )
                    st.session_state.data.properties.append(new_property)
                    st.success("✅ Property added successfully!")
                    st.rerun()
                else:
                    st.error("Please fill all required fields")
        
        st.subheader("Current Properties")
        if st.session_state.data.properties:
            for idx, prop in enumerate(st.session_state.data.properties):
                with st.expander(f"Property {idx+1}: {prop.property_type}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Owner:** {prop.owner_name}")
                        st.write(f"**Type:** {prop.property_type}")
                        st.write(f"**Address:** {prop.address}")
                    with col2:
                        st.write(f"**Valuation (INR):** ₹{prop.valuation_inr:,.2f}")
                        st.write(f"**Valuation ({st.session_state.data.foreign_currency}):** {prop.valuation_foreign:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_prop_{idx}"):
                            st.session_state.data.properties.pop(idx)
                            st.rerun()
            
            st.metric("Total Property Value (INR)", f"₹{st.session_state.data.total_immovable_assets_inr:,.2f}")
        else:
            st.info("No properties added yet")
    
    # Tab 14: Liabilities
    with tabs[13]:
        st.header("Liabilities")
        
        st.subheader("Add Liability")
        with st.form("liability_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            with col1:
                description = st.text_input("Liability Description (e.g., Home Loan, Personal Loan)")
                amount = st.number_input("Amount (INR)", min_value=0.0, step=10000.0, format="%.2f")
            with col2:
                details = st.text_area("Additional Details", height=100)
            
            if st.form_submit_button("➕ Add Liability"):
                if description:
                    new_liability = Liability(
                        description=description,
                        amount_inr=amount,
                        details=details
                    )
                    st.session_state.data.liabilities.append(new_liability)
                    st.success("✅ Liability added successfully!")
                    st.rerun()
                else:
                    st.error("Please enter liability description")
        
        st.subheader("Current Liabilities")
        if st.session_state.data.liabilities:
            for idx, liab in enumerate(st.session_state.data.liabilities):
                with st.expander(f"Liability {idx+1}: {liab.description}"):
                    col1, col2, col3 = st.columns([3, 2, 1])
                    with col1:
                        st.write(f"**Description:** {liab.description}")
                        st.write(f"**Details:** {liab.details}")
                    with col2:
                        st.write(f"**Amount (INR):** ₹{liab.amount_inr:,.2f}")
                    with col3:
                        if st.button("🗑️ Delete", key=f"del_liab_{idx}"):
                            st.session_state.data.liabilities.pop(idx)
                            st.rerun()
            
            st.metric("Total Liabilities (INR)", f"₹{st.session_state.data.total_liabilities_inr:,.2f}")
        else:
            st.info("No liabilities added - Net Worth will equal Total Assets")
    
    # Tab 15: Summary & Generate
    with tabs[14]:
        st.header("Summary & Generate Certificate")
        
        # Summary Cards
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Movable Assets", f"₹{st.session_state.data.total_movable_assets_inr:,.2f}")
            st.caption(f"Bank: ₹{st.session_state.data.total_bank_balance_inr:,.2f}")
            st.caption(f"Insurance: ₹{st.session_state.data.total_insurance_inr:,.2f}")
            st.caption(f"P.F. Accounts: ₹{st.session_state.data.total_pf_accounts_inr:,.2f}")
            st.caption(f"Deposits: ₹{st.session_state.data.total_deposits_inr:,.2f}")
            st.caption(f"NPS: ₹{st.session_state.data.total_nps_inr:,.2f}")
            st.caption(f"Mutual Funds: ₹{st.session_state.data.total_mutual_funds_inr:,.2f}")
            st.caption(f"Shares: ₹{st.session_state.data.total_shares_inr:,.2f}")
            st.caption(f"Vehicles: ₹{st.session_state.data.total_vehicles_inr:,.2f}")
            st.caption(f"Post Office: ₹{st.session_state.data.total_post_office_inr:,.2f}")
            st.caption(f"Partnership Firms: ₹{st.session_state.data.total_partnership_firms_inr:,.2f}")
            st.caption(f"Gold: ₹{st.session_state.data.total_gold_inr:,.2f}")
        
        with col2:
            st.metric("Total Immovable Assets", f"₹{st.session_state.data.total_immovable_assets_inr:,.2f}")
            st.caption(f"Properties: {len(st.session_state.data.properties)}")
        
        with col3:
            st.metric("Total Liabilities", f"₹{st.session_state.data.total_liabilities_inr:,.2f}")
        
        st.markdown("---")
        
        # Net Worth Display
        st.subheader("💰 NET WORTH CALCULATION")
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric(
                "Net Worth (INR)", 
                f"₹{st.session_state.data.net_worth_inr:,.2f}",
                help="Total Assets - Total Liabilities"
            )
        
        with col2:
            st.metric(
                f"Net Worth ({st.session_state.data.foreign_currency})", 
                f"{st.session_state.data.net_worth_foreign:,.2f}",
                help=f"At exchange rate: {st.session_state.data.exchange_rate} INR = 1 {st.session_state.data.foreign_currency}"
            )
        
        st.markdown("---")
        
        # Validation
        validation_errors = []
        if not st.session_state.data.individual_name:
            validation_errors.append("❌ Individual's name is required")
        if not st.session_state.data.individual_address:
            validation_errors.append("❌ Individual's address is required")
        if not st.session_state.data.embassy_name:
            validation_errors.append("❌ Embassy/Consulate name is required")
        if len(st.session_state.data.bank_accounts) == 0 and len(st.session_state.data.properties) == 0:
            validation_errors.append("⚠️ Add at least one bank account or property")
        
        if validation_errors:
            st.error("Please complete the following:")
            for error in validation_errors:
                st.write(error)
        else:
            st.success("✅ All required information provided. Ready to generate certificate!")
        
        # Generate Button
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("📄 Generate Net Worth Certificate", type="primary", use_container_width=True, disabled=bool(validation_errors)):
                with st.spinner("Generating certificate..."):
                    try:
                        # Generate document
                        doc = generate_networth_certificate(st.session_state.data)
                        
                        # Save to bytes
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        # Download button
                        st.success("✅ Certificate generated successfully!")
                        st.download_button(
                            label="⬇️ Download Certificate (DOCX)",
                            data=doc_io.getvalue(),
                            file_name=f"NetWorth_Certificate_{st.session_state.data.individual_name.replace(' ', '_')}_{datetime.date.today().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error(f"Error generating certificate: {str(e)}")
                        st.exception(e)

if __name__ == "__main__":
    main()
