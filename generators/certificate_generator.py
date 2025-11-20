"""
Main certificate generation module
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from models import NetWorthData
from generators.annexure_generator import generate_annexures
from config import DEFAULT_FONT_NAME, DEFAULT_FONT_SIZE_PT


def _set_document_margins(doc, left=0.75, right=0.75, top=1.0, bottom=1.0):
    """
    Set consistent margins for all sections in the document.
    
    Args:
        doc: Document object
        left: Left margin in inches (default: 0.75)
        right: Right margin in inches (default: 0.75)
        top: Top margin in inches (default: 1.0)
        bottom: Bottom margin in inches (default: 1.0)
    """
    for section in doc.sections:
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)


def _create_signature_table(doc, data: NetWorthData):
    """Create signature table with invisible borders"""
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = True

    # Make table borders invisible using XML
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                  r'<w:top w:val="nil"/>'
                                  r'<w:left w:val="nil"/>'
                                  r'<w:bottom w:val="nil"/>'
                                  r'<w:right w:val="nil"/>'
                                  r'</w:tcBorders>')
            tcPr.append(tcBorders)

    # Left column - CA details
    sig_table.rows[0].cells[0].text = f'{data.ca_partner_name.upper()}'
    sig_table.rows[1].cells[0].text = f'{data.ca_designation.upper()}'
    sig_table.rows[2].cells[0].text = f'MEMBERSHIP NO.: {data.ca_membership_no}'
    sig_table.rows[3].cells[0].text = 'UDIN: [TO BE GENERATED]'

    # Right column - Date and Place
    sig_table.rows[0].cells[1].text = f'DATE: {data.certificate_date}'
    sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sig_table.rows[1].cells[1].text = f'PLACE: {data.ca_place.upper()}'
    sig_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Leave bottom two cells in right column empty for spacing
    sig_table.rows[2].cells[1].text = ''
    sig_table.rows[3].cells[1].text = ''


def generate_networth_certificate(data: NetWorthData) -> Document:
    """
    Generate the complete Net Worth Certificate document
    
    Args:
        data: NetWorthData object containing all certificate information
        
    Returns:
        Document object ready to be saved
    """
    doc = Document()
    
    # Set narrower margins to accommodate longer title lines
    # Default margins are 1 inch; reducing left/right to 0.75 inches for more width
    _set_document_margins(doc, left=0.75, right=0.75, top=1.0, bottom=1.0)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = Pt(DEFAULT_FONT_SIZE_PT)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Independent Practitioner's Certificate on Net Worth where no Books of")
    title_run.bold = True
    title_run.underline = True
    title_run.add_break()
    title_run = title.add_run("Account have been maintained (For VISA Application Purpose)")
    title_run.bold = True
    title_run.underline = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # To Address
    to_para = doc.add_paragraph('To')
    to_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    embassy_para = doc.add_paragraph(data.embassy_name)
    embassy_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    embassy_lines = data.embassy_address.split('\n')
    for line in embassy_lines:
        line_para = doc.add_paragraph(line)
        line_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Certificate Body
    para = doc.add_paragraph()
    run = para.add_run("Independent Practitioner's Certificate on Net Worth where no Books of")
    run.underline = True
    run.add_break()
    run = para.add_run("Account have been maintained (For VISA Application Purpose)")
    run.underline = True
    
    cert_para1 = doc.add_paragraph(
        f'1. This Certificate is issued in accordance with the terms of my/our engagement letter/agreement dated {data.engagement_date}.'
    )
    cert_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Helper to describe individual(s) with passports
    if data.individuals:
        primary = data.individuals[0]
        primary_name_with_passport = (
            f"{primary.full_name} (Passport No.: {primary.passport_number})"
            if primary.passport_number
            else primary.full_name
        )
        if len(data.individuals) == 1:
            individuals_phrase = primary_name_with_passport
        else:
            other_count = len(data.individuals) - 1
            individuals_phrase = f"{primary_name_with_passport} and {other_count} other individual(s)"
    else:
        primary_name_with_passport = "Unnamed Individual"
        individuals_phrase = primary_name_with_passport
    # Build combined address string (may be same or different for each individual)
    if data.individuals:
        unique_addresses = {
            ind.address for ind in data.individuals if ind.address.strip()
        }
        if unique_addresses:
            if len(unique_addresses) == 1:
                address_text = next(iter(unique_addresses))
            else:
                address_text = "; ".join(unique_addresses)
        else:
            address_text = ""
    else:
        address_text = ""

    cert_para2 = doc.add_paragraph(
        f'2. I/we have been engaged by {individuals_phrase} (hereinafter referred to as the "individuals") '
        f'having residential address(es) at {address_text} to certify the Net Worth as at {data.certificate_date} '
        f'for submission to {data.embassy_name} for VISA application purpose.'
    )
    cert_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Individual's Responsibility
    para = doc.add_paragraph()
    run = para.add_run("Individual's Responsibility")
    run.bold = True
    run.underline = True
    
    resp_para = doc.add_paragraph(
        f'3. The individual is responsible for preparing the Statement of Net Worth ("the Statement") as at {data.certificate_date} '
        f'and for maintaining adequate records and internal controls to support the accuracy and completeness of the information contained therein.'
    )
    resp_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Practitioner's Responsibility
    para = doc.add_paragraph()
    run = para.add_run("Practitioner's Responsibility")
    run.bold = True
    run.underline = True
    
    prac_para = doc.add_paragraph(
        f'4. My/our responsibility is to examine and certify the Statement of Net Worth as at {data.certificate_date} '
        f'based on the supporting documents provided. The examination was performed in accordance with the ICAI Guidance Note '
        f'on Reports or Certificates for Special Purposes, and in compliance with the ICAI Code of Ethics. '
        f'I/we have also followed the relevant requirements of SQC 1 relating to quality control.'
    )
    prac_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Opinion
    para = doc.add_paragraph()
    run = para.add_run("Opinion")
    run.bold = True
    run.underline = True
    
    opinion_para = doc.add_paragraph(
        f'7. On the basis of the examination carried out and the information and explanations furnished to me/us, '
        f'I/we certify that the annexed Statement of Net Worth of {individuals_phrase} as at {data.certificate_date} '
        f'presents a Net Worth of ₹{data.net_worth_inr:,.2f}, derived from the records, representations and supporting documents '
        f'provided by the individual.'
    )
    opinion_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Restriction on Use
    para = doc.add_paragraph()
    run = para.add_run("Restriction on Use")
    run.bold = True
    run.underline = True
    
    restrict_para = doc.add_paragraph(
        f'8. This Certificate is prepared at the individual\'s request for submission to {data.embassy_name} for VISA processing. '
        f'It is restricted to this purpose only and is not intended for any other use. No responsibility or liability is accepted '
        f'towards any person other than the specified addressee without my/our written consent.'
    )
    restrict_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature Block
    firm_para = doc.add_paragraph(f'FOR {data.ca_firm_name.upper()}')
    firm_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ca_para = doc.add_paragraph('CHARTERED ACCOUNTANTS')
    ca_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    frn_para = doc.add_paragraph(f"FRN: {data.ca_frn}")
    frn_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    _create_signature_table(doc, data)
    
    doc.add_paragraph()

    enclosure_para = doc.add_paragraph(
        f"Enclosure: Statement of Net Worth of {individuals_phrase} as at {data.certificate_date}"
    )
    enclosure_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add page break for annexures
    doc.add_page_break()
    
    # Generate annexures
    generate_annexures(doc, data)
    
    return doc

