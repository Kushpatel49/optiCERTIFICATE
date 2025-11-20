"""
Annexure generation module for Net Worth Certificate
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from models import NetWorthData
from generators.table_utils import add_table_with_borders, enforce_sr_no_column_width

def generate_annexures(doc, data: NetWorthData):
    """Generate all annexures"""
    
    # Annexure Header
    title = doc.add_paragraph()
    title_run = title.add_run('Annexure – Statement of Net Worth')
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()

    # Build individuals line(s)
    if data.individuals:
        if len(data.individuals) == 1:
            ind = data.individuals[0]
            name_with_passport = (
                f"{ind.full_name} (Passport No.: {ind.passport_number})"
                if ind.passport_number
                else ind.full_name
            )
            name_para = doc.add_paragraph(f"Name of Individual: {name_with_passport}")
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            address_para = doc.add_paragraph(f"Address: {ind.address}")
            address_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            # Multiple individuals – list each on its own line
            name_para = doc.add_paragraph("Individuals:")
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for ind in data.individuals:
                name_with_passport = (
                    f"{ind.full_name} (Passport No.: {ind.passport_number})"
                    if ind.passport_number
                    else ind.full_name
                )
                line_para = doc.add_paragraph(f" - {name_with_passport}")
                line_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            address_para = doc.add_paragraph("Addresses:")
            address_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            unique_addresses = {
                ind.address for ind in data.individuals if ind.address.strip()
            }
            for addr in unique_addresses:
                line_para = doc.add_paragraph(f" - {addr}")
                line_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    date_para = doc.add_paragraph(f'Date of Certificate: {data.certificate_date}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    purpose_para = doc.add_paragraph(f'Purpose: VISA Application – Submission to {data.embassy_name}')
    purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run(f'Statement of Net Worth as at {data.certificate_date}')
    run.bold = True
    
    doc.add_paragraph()
    
    # Summary - Net Worth Table
    summary_title = doc.add_paragraph()
    summary_title_run = summary_title.add_run('SUMMARY - NET WORTH')
    summary_title_run.bold = True
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Create summary table: 4 rows (header + 3 data rows + 1 total row) = 5 rows, 4 columns
    summary_table = add_table_with_borders(doc, 5, 4)
    
    # Header row
    summary_headers = [
        'Particulars',
        'Estimated Market Value (INR)',
        f'Estimated Market Value ({data.foreign_currency}@ {data.exchange_rate})',
        'Annexure',
    ]
    for i, header in enumerate(summary_headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Apply standardized column-width logic so the first column and numeric
    # columns have enough width to avoid wrapping
    enforce_sr_no_column_width(summary_table, summary_headers)
    
    # Row 1: Movable Assets
    summary_table.rows[1].cells[0].text = 'Movable Assets'
    summary_table.rows[1].cells[1].text = f'{data.total_movable_assets_inr:,.2f}'
    summary_table.rows[1].cells[2].text = f'{data.total_movable_assets_foreign:,.2f}'
    summary_table.rows[1].cells[3].text = '(i)'
    
    # Row 2: Immovable Assets
    summary_table.rows[2].cells[0].text = 'Immovable Assets'
    summary_table.rows[2].cells[1].text = f'{data.total_immovable_assets_inr:,.2f}'
    summary_table.rows[2].cells[2].text = f'{data.total_immovable_assets_foreign:,.2f}'
    summary_table.rows[2].cells[3].text = '(ii)'
    
    # Row 3: Liabilities
    summary_table.rows[3].cells[0].text = 'Liabilities'
    if data.total_liabilities_inr > 0:
        summary_table.rows[3].cells[1].text = f'{data.total_liabilities_inr:,.2f}'
        summary_table.rows[3].cells[2].text = f'{data.total_liabilities_foreign:,.2f}'
    else:
        summary_table.rows[3].cells[1].text = '-'
        summary_table.rows[3].cells[2].text = '-'
    summary_table.rows[3].cells[3].text = '(iii)'
    
    # Row 4: Total (i+ii-iii)
    summary_table.rows[4].cells[0].text = 'Total (i+ii-iii)'
    summary_table.rows[4].cells[0].paragraphs[0].runs[0].bold = True
    summary_table.rows[4].cells[1].text = f'{data.net_worth_inr:,.2f}'
    summary_table.rows[4].cells[1].paragraphs[0].runs[0].bold = True
    summary_table.rows[4].cells[2].text = f'{data.net_worth_foreign:,.2f}'
    summary_table.rows[4].cells[2].paragraphs[0].runs[0].bold = True
    summary_table.rows[4].cells[3].text = ''
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Annexure (i) - Movable Assets Summary
    para = doc.add_paragraph()
    run = para.add_run('Annexure (i) - Movable Assets')
    run.bold = True
    
    # Build list of categories with data
    categories = []
    if data.bank_accounts:
        categories.append(('1', 'Bank Account', 'A', data.total_bank_balance_inr, data.total_bank_balance_foreign))
    if data.insurance_policies:
        categories.append(('2', 'LIC', 'B', data.total_insurance_inr, data.total_insurance_foreign))
    if data.pf_accounts:
        categories.append(('3', 'P.F. Account', 'C', data.total_pf_accounts_inr, data.total_pf_accounts_foreign))
    if data.deposits:
        categories.append(('4', 'Deposit', 'D', data.total_deposits_inr, data.total_deposits_foreign))
    if data.nps_accounts:
        categories.append(('5', 'NPS', 'E', data.total_nps_inr, data.total_nps_foreign))
    if data.mutual_funds:
        categories.append(('6', 'Investment in Mutual Fund', 'F', data.total_mutual_funds_inr, data.total_mutual_funds_foreign))
    if data.shares:
        categories.append(('7', 'Shares & Securities', 'G', data.total_shares_inr, data.total_shares_foreign))
    if data.vehicles:
        categories.append(('8', 'Vehicles', 'H', data.total_vehicles_inr, data.total_vehicles_foreign))
    if data.post_office_schemes:
        categories.append(('9', 'Post Office Schemes', 'I', data.total_post_office_inr, data.total_post_office_foreign))
    if data.partnership_firms:
        categories.append(('10', 'Investments in Partnership Firms', 'J', data.total_partnership_firms_inr, data.total_partnership_firms_foreign))
    if data.gold_holdings:
        categories.append(('11', 'Gold', 'K', data.total_gold_inr, data.total_gold_foreign))
    
    # Only create table if there are categories with data
    if categories:
        # Create table with header + data rows + total row
        table = add_table_with_borders(doc, len(categories) + 2, 5)
    
    # Header row
    headers = ['Sr. No.', 'Particulars', 'Sub-Annexure', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if header == 'Sr. No.':
            cell.text = 'Sr.\nNo.'
        else:
            cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
        enforce_sr_no_column_width(table, headers)
        
        # Data rows - only for categories with data
        for idx, (sr_no, particular, sub_annexure, inr_amount, foreign_amount) in enumerate(categories, 1):
            table.rows[idx].cells[0].text = sr_no
            table.rows[idx].cells[1].text = particular
            table.rows[idx].cells[2].text = sub_annexure
            table.rows[idx].cells[3].text = f'{inr_amount:,.2f}'
            table.rows[idx].cells[4].text = f'{foreign_amount:,.2f}'
        
        # Total row
        total_row = len(categories) + 1
        table.rows[total_row].cells[0].text = ''
        cell = table.rows[total_row].cells[1]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        table.rows[total_row].cells[2].text = ''
        cell = table.rows[total_row].cells[3]
        cell.text = f'{data.total_movable_assets_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[total_row].cells[4]
        cell.text = f'{data.total_movable_assets_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    else:
        # If no movable assets, add a note
        doc.add_paragraph('No movable assets to report.')
    
    # Sub Annexure A - Bank Accounts
    if data.bank_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (A) – Bank Account')
        run.bold = True
        table = add_table_with_borders(doc, len(data.bank_accounts) + 2, 6)
        
        # Headers
        headers = ['Sr. No.', 'Name of the Account Holder', 'Account No.', 'Bank Name', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        # Data rows
        for idx, acc in enumerate(data.bank_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = acc.holder_name
            table.rows[idx].cells[2].text = acc.account_number
            table.rows[idx].cells[3].text = acc.bank_name
            table.rows[idx].cells[4].text = f'{acc.balance_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{acc.balance_foreign:,.2f}'
        
        # Total row - keep Sr. No. column empty so "Total" does not wrap
        last_row = len(data.bank_accounts) + 1
        table.rows[last_row].cells[0].text = ''
        cell = table.rows[last_row].cells[1]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_bank_balance_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_bank_balance_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.bank_accounts_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.bank_accounts_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Sub Annexure B - Insurance Policies
    if data.insurance_policies:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (B): LIFE INSURANCE POLICIES')
        run.bold = True
        table = add_table_with_borders(doc, len(data.insurance_policies) + 2, 5)
        
        headers = ['Sr. No.', 'POLICY Holder', 'Policy No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, policy in enumerate(data.insurance_policies, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = policy.holder_name
            table.rows[idx].cells[2].text = policy.policy_number
            table.rows[idx].cells[3].text = f'{policy.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{policy.amount_foreign:,.2f}'
        
        last_row = len(data.insurance_policies) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'TOTAL'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_insurance_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_insurance_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
    
        # Add notes if provided
        if data.insurance_policies_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.insurance_policies_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Sub Annexure C - P.F. Accounts
    if data.pf_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (C) - P.F Account')
        run.bold = True
        table = add_table_with_borders(doc, len(data.pf_accounts) + 2, 5)
        headers = ['Sr. No.', 'Name of the Account Holder', 'PF Account No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, acc in enumerate(data.pf_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = acc.holder_name
            table.rows[idx].cells[2].text = acc.pf_account_number
            table.rows[idx].cells[3].text = f'{acc.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{acc.amount_foreign:,.2f}'
        
        last_row = len(data.pf_accounts) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_pf_accounts_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_pf_accounts_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.pf_accounts_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.pf_accounts_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sub Annexure D - Deposits
    if data.deposits:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (D): Deposit')
        run.bold = True
        table = add_table_with_borders(doc, len(data.deposits) + 2, 5)
        headers = ['Sr. No.', 'Name of Investment Holder', 'A/C Number', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, dep in enumerate(data.deposits, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = dep.holder_name
            table.rows[idx].cells[2].text = dep.account_number
            table.rows[idx].cells[3].text = f'{dep.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{dep.amount_foreign:,.2f}'
        
        last_row = len(data.deposits) + 1
        # Leave Sr. No. column empty; place "Total" in the wider second column
        table.rows[last_row].cells[0].text = ''
        cell = table.rows[last_row].cells[1]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_deposits_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_deposits_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.deposits_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.deposits_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Sub Annexure E - NPS
    if data.nps_accounts:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (E) - NPS')
        run.bold = True
        table = add_table_with_borders(doc, len(data.nps_accounts) + 2, 5)
        headers = ['Sr. No.', 'Name of Owner', 'PRAN No.', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, nps in enumerate(data.nps_accounts, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = nps.owner_name
            table.rows[idx].cells[2].text = nps.pran_number
            table.rows[idx].cells[3].text = f'{nps.amount_inr:,.2f}'
            table.rows[idx].cells[4].text = f'{nps.amount_foreign:,.2f}'
        
        last_row = len(data.nps_accounts) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_nps_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_nps_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.nps_accounts_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.nps_accounts_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sub Annexure F - Investment in Mutual Fund
    if data.mutual_funds:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (F) - Investment in Mutual Fund')
        run.bold = True
        table = add_table_with_borders(doc, len(data.mutual_funds) + 2, 6)
        headers = ['Sr. No.', 'Name of the Account Holder', 'Policy/Folio Number', 'Policy Name', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)

        for idx, mf in enumerate(data.mutual_funds, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = mf.holder_name
            table.rows[idx].cells[2].text = mf.folio_number
            table.rows[idx].cells[3].text = mf.policy_name
            table.rows[idx].cells[4].text = f'{mf.amount_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{mf.amount_foreign:,.2f}'

        last_row = len(data.mutual_funds) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_mutual_funds_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_mutual_funds_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.mutual_funds_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.mutual_funds_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sub Annexure G - Gold
    if data.gold_holdings:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Sub Annexure (G) - Gold')
        run.bold = True
        table = add_table_with_borders(doc, len(data.gold_holdings) + 2, 6)
        
        headers = ['Sr. No.', 'Name of Party', 'Weight (gram)', 'Rate/10 g (Rs.)', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, gold in enumerate(data.gold_holdings, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = gold.owner_name
            table.rows[idx].cells[2].text = f'{gold.weight_grams:.3f}'
            table.rows[idx].cells[3].text = f'{gold.rate_per_10g:,.2f}'
            table.rows[idx].cells[4].text = f'{gold.amount_inr:,.2f}'
            table.rows[idx].cells[5].text = f'{gold.amount_foreign:,.2f}'
        
        last_row = len(data.gold_holdings) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[4]
        cell.text = f'{data.total_gold_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[5]
        cell.text = f'{data.total_gold_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        if data.gold_holdings and data.gold_holdings[0].valuation_date:
            valuer_para = doc.add_paragraph(f'As per the Property Valuation Certificates dated {data.gold_holdings[0].valuation_date} issued by Approved Valuer {data.gold_holdings[0].valuer_name}')
            valuer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add notes if provided
        if data.gold_holdings_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.gold_holdings_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Annexure (ii) - Immovable Assets
    if data.properties:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Annexure (ii) - Immovable Assets')
        run.bold = True
        table = add_table_with_borders(doc, len(data.properties) + 2, 4)
        
        headers = ['Sr. No.', 'Particulars of Property', 'Amount in INR', f'Amount in {data.foreign_currency}@ {data.exchange_rate} INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, prop in enumerate(data.properties, 1):
            table.rows[idx].cells[0].text = str(idx)
            property_text = f'{prop.owner_name}\n\n{prop.property_type}\n\n{prop.address}'
            if prop.valuation_date and prop.valuer_name:
                property_text += f'\n\n(Valuation as on {prop.valuation_date} by {prop.valuer_name})'
            table.rows[idx].cells[1].text = property_text
            table.rows[idx].cells[2].text = f'{prop.valuation_inr:,.2f}'
            table.rows[idx].cells[3].text = f'{prop.valuation_foreign:,.2f}'
        
        last_row = len(data.properties) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[2]
        cell.text = f'{data.total_immovable_assets_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_immovable_assets_foreign:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.properties_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.properties_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Annexure (iii) - Liabilities
    if data.liabilities:
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run('Annexure (iii) - Liabilities')
        run.bold = True
        table = add_table_with_borders(doc, len(data.liabilities) + 2, 4)
        
        headers = ['Sr. No.', 'Description', 'Details', 'Amount in INR']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            if header == 'Sr. No.':
                cell.text = 'Sr.\nNo.'
            else:
                cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        enforce_sr_no_column_width(table, headers)
        
        for idx, liab in enumerate(data.liabilities, 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = liab.description
            table.rows[idx].cells[2].text = liab.details
            table.rows[idx].cells[3].text = f'{liab.amount_inr:,.2f}'
        
        last_row = len(data.liabilities) + 1
        cell = table.rows[last_row].cells[0]
        cell.text = 'Total'
        cell.paragraphs[0].runs[0].bold = True
        cell = table.rows[last_row].cells[3]
        cell.text = f'{data.total_liabilities_inr:,.2f}'
        cell.paragraphs[0].runs[0].bold = True
        
        # Add notes if provided
        if data.liabilities_notes:
            doc.add_paragraph()
            notes_para = doc.add_paragraph(f'Notes: {data.liabilities_notes}')
            notes_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Net Worth Calculation
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run(f'Net Worth: ₹{data.net_worth_inr:,.2f}')
    run.bold = True
    run.font.size = Pt(14)
    
    # Notes
    doc.add_paragraph()

    notes_title = doc.add_paragraph('Notes:')
    notes_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note1 = doc.add_paragraph('1. The above Statement is prepared based on details and supporting documents provided by the individual.')
    note1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note2 = doc.add_paragraph('2. Valuation of assets is based on self-declaration / available records and has not been independently verified unless specified.')
    note2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note3 = doc.add_paragraph(f'3. This Annexure should be read with the Certificate dated {data.certificate_date} issued by the undersigned.')
    note3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note4 = doc.add_paragraph('4. Loan documents and related confirmations were not made available for verification. As informed to us there is no any liability as on the date.')
    note4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    note5 = doc.add_paragraph('5. The Information Furnished in the Certificate do not certify any Title Neither Ownership as we are not Legal Expert.')
    note5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Final Signature
    final_firm_para = doc.add_paragraph(f'FOR, {data.ca_firm_name.upper()}')
    final_firm_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    final_ca_para = doc.add_paragraph('CHARTERED ACCOUNTANTS')
    final_ca_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    final_frn_para = doc.add_paragraph(f'FRN: {data.ca_frn}')
    final_frn_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()
    doc.add_paragraph()

    # Create two-column layout for signature
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = True

    # Make table borders invisible using XML
    for row in sig_table.rows:
        for cell in row.cells:
            # Set all borders to nil (invisible)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                  r'<w:top w:val="nil"/>'
                                  r'<w:left w:val="nil"/>'
                                  r'<w:bottom w:val="nil"/>'
                                  r'<w:right w:val="nil"/>'
                                  r'</w:tcBorders>')
            tcPr.append(tcBorders)

    # Left column - CA details
    sig_table.rows[0].cells[0].text = f'{data.ca_partner_name.upper()}'
    sig_table.rows[1].cells[0].text = f'{data.ca_designation.upper()}'
    sig_table.rows[2].cells[0].text = f'MEMBERSHIP NO.: {data.ca_membership_no}'
    sig_table.rows[3].cells[0].text = 'UDIN: [TO BE GENERATED]'

    # Right column - Date and Place
    sig_table.rows[0].cells[1].text = f'DATE: {data.certificate_date}'
    sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sig_table.rows[1].cells[1].text = f'PLACE: {data.ca_place.upper()}'
    sig_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Leave bottom two cells in right column empty for spacing
    sig_table.rows[2].cells[1].text = ''
    sig_table.rows[3].cells[1].text = ''

